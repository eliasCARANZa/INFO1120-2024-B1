import sqlite3,pandas as pd,matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Cm, Mm

def obtener_datos_desde_bd(bd):
    conexion = sqlite3.connect(bd)
    consulta = '''                   
    SELECT personas.*, Salarios.*
    FROM personas
    INNER JOIN Salarios ON personas.id_rol = id_Salarios
    '''
    df = pd.read_sql_query(consulta, conexion)
    conexion.close()
    return df

def example_contract(date: str, rol: str, address: str, rut: str, full_name: str, nationality: str, birth_date: str, profession: str, salary: str):
    document = Document()
    font = document.styles['Normal'].font
    font.name = 'Book Antiqua'
    sections = document.sections
    for section in sections:
        section.page_height = Cm(35.56)
        section.page_width = Cm(21.59)
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
        section.header_distance = Mm(12.7)
        section.footer_distance = Mm(12.7)

    header = document.sections[0].header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture("imagenes/header.png")
    document.add_picture('imagenes/logo.png')
    h = document.add_paragraph('')
    h.add_run('CONTRATO DE PRESTACIÓN DE SERVICIOS A HONORARIOS\n').bold = True
    h.alignment = 1
    font.size = Pt(10)
    h.add_run(
        '............................................................................................................................................................................................................................. ').bold = True

    p = document.add_paragraph(f'En Temuco, a {str(date)}, entre la Corporación de Innovación y Desarrollo Tecnológico, Rut 78.898.766-4, representada por su Director General don(a) Roberto Gomez Bolainas, Cédula de Identidad Nº 10.678.990-2, ambos domiciliados en Caupolican 455 de esta ciudad, en adelante la “Corporación” y  {full_name}, de nacionalidad {nationality}, de profesión {profession}, nacido el {birth_date}, con domicilio en {address}, Cédula de Identidad N° {rut}, en adelante, el “Prestador de Servicios”, se ha convenido el siguiente contrato de prestación de servicios a honorarios: \n')
    font.size = Pt(8)
    p1 = document.add_paragraph('')
    p1.add_run('PRIMERO        :').bold = True
    p1.add_run('En el marco del acuerdo de servicios profesionales fechado el 11 de noviembre de 2020, establecido entre la Agencia Nacional de Estándares Educativos y la Corporación de Innovación y Desarrollo Tecnológico , y ratificado según la Resolución Exenta N°603 del 23 de noviembre de 2020, la Corporación encarga los servicios profesionales del Prestador de Servicios, para que ejecute la siguiente tarea en el proyecto "Evaluación de competencias específicas y metodologías de aprendizaje artificial 2020, ID 67703-20-JJ90."')
    p1.add_run('\n SEGUNDO        :').bold = True
    p1.add_run('El rol a desempeñar es de '+rol+'.')
    p1.add_run('\n TERCERO        :').bold = True
    p1.add_run('El plazo para la realización de la prestación de servicios encomendada será el '+str(date))
    p1.add_run('\n CUARTO        :').bold = True
    p1.add_run('Por el servicio profesional efectivamente realizado, se pagara un monto bruto variable, el cual corresponderá a cada rol dentro de la empresa capacitación, de acuerdo al siguiente detalle: ')
    table = document.add_table(rows=2, cols=2)
    table.alignment = 1
    hdr_cells0 = table.rows[0].cells
    hdr_cells0[0].text='Rol'
    hdr_cells0[1].text='Monto Bruto'
    hdr_cells = table.rows[1].cells
    hdr_cells[0].text = rol
    hdr_cells[1].text= salary
    p1.add_run('\n QUINTO        :').bold = True
    p1.add_run('El Prestador de Servicios acepta el encargo y las condiciones precedentes.')
    p1.add_run('\n SEXTO        :').bold = True
    p1.add_run('El Prestador de Servicios está obligado a mantener la confidencialidad de todos los materiales utilizados, conforme al Acuerdo de Confidencialidad previamente establecido.')
    p1.add_run('\n En comprobante, previa lectura y ratificación, las partes firman.  ').bold = True
    table = document.add_table(rows=2, cols=2)
    table.alignment = 1
    hdr_cells0 = table.rows[0].cells[1].add_paragraph()
    r = hdr_cells0.add_run()
    r.add_picture('imagenes/firma.png')
    hdr_cells = table.rows[1].cells
    hdr_cells[0].text = '-----------------------------------------------------------\nEL PRESTADOR DE SERVICIOS'
    hdr_cells[1].text = '-----------------------------------------------------------\np. LA CORPORACION'
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    run = paragraph.add_run('Caupolican 0455, Temuco, Chile, www.corpoindet.cl')
    run.add_picture("imagenes/footer1.png")
    document.save(f'{full_name}.docx')

def grafico_promedio_sueldo_profesion(df):
    sueldo_promedio_profesion = df.groupby('profesion')['Sueldo'].mean()
    plt.figure(figsize=(10, 6))
    sueldo_promedio_profesion.plot(kind='bar', color='skyblue')
    plt.xlabel('Profesión')
    plt.ylabel('Promedio de Sueldo')
    plt.title('Promedio de Sueldo por Profesión')
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()
    plt.savefig('grafico_promedio_sueldo_profesion.png')
    plt.show()

def grafico_distribucion_profesiones(df):
    profesion_distribucion = df['profesion'].value_counts()
    plt.figure(figsize=(8, 8))
    profesion_distribucion.plot(kind='pie', autopct='%1.1f%%', startangle=140)
    plt.title('Distribución de Profesiones')
    plt.ylabel('')
    plt.savefig('grafico_distribucion_profesiones.png')
    plt.show()

def grafico_conteo_profesionales_nacionalidad(df):
    nacionalidad_conteo = df['nacionalidad'].value_counts()
    plt.figure(figsize=(10, 6))
    nacionalidad_conteo.plot(kind='bar', color='salmon')
    plt.xlabel('Nacionalidad')
    plt.ylabel('Conteo de Profesionales')
    plt.title('Conteo de Profesionales por Nacionalidad')
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()
    plt.savefig('grafico_conteo_profesionales_nacionalidad.png')
    plt.show()

bd = 'db_personas.db'
df = obtener_datos_desde_bd(bd)

def menu():
    print("Seleccione la acción que desea realizar:")
    print("1. Generar contratos")
    print("2. Visualizar gráficos")
    print("3. Salir")

def menu_contratos():
    print("Seleccione el tipo de generación de contratos:")
    print("1. Generar contrato para una persona")
    print("2. Generar contratos para un rango de personas")
    print("3. Volver al menú principal")

def menu_graficos():
    print("Seleccione el gráfico que desea visualizar:")
    print("1. Gráfico promedio sueldo por profesión")
    print("2. Gráfico de distribución de profesiones")
    print("3. Gráfico de conteo de profesionales por nacionalidad")
    print("4. Volver al menú principal")

opciones_graficos = {
    1: grafico_promedio_sueldo_profesion,
    2: grafico_distribucion_profesiones,
    3: grafico_conteo_profesionales_nacionalidad
}

print(df)

while True:
    menu()
    try:
        opcion = int(input("Ingrese el número de la opción deseada: "))
        if opcion == 1:
            while True:
                menu_contratos()
                try:
                    sub_opcion = int(input("Ingrese el número de la opción deseada: "))
                    if sub_opcion == 1:
                        indice = int(input("Ingrese el índice de la persona para generar el contrato: "))
                        if indice < 0 or indice >= len(df):
                            raise IndexError("Índice fuera de rango")

                        persona_seleccionada = df.iloc[indice]
                        date = persona_seleccionada['fecha_ingreso']
                        rol = persona_seleccionada['Rol']
                        address = persona_seleccionada['residencia']
                        rut = persona_seleccionada['rut']
                        full_name = persona_seleccionada['nombre_completo']
                        nationality = persona_seleccionada['nacionalidad']
                        birth_date = persona_seleccionada['fecha_de_nacimiento']
                        profession = persona_seleccionada['profesion']
                        salary = persona_seleccionada['Sueldo']

                        example_contract(date, rol, address, rut, full_name, nationality, birth_date, profession, str(salary))
                        print(f"Contrato generado para {full_name}.")

                    elif sub_opcion == 2:
                        inicio = int(input("Ingrese el índice de inicio para generar los contratos: "))
                        fin = int(input("Ingrese el índice final para generar los contratos: "))

                        if (inicio < 0 or inicio >= len(df)) or (fin < 0 or fin >= len(df)):
                            raise IndexError("Índice fuera de rango")

                        if inicio > fin:
                            raise ValueError("El índice de inicio debe ser menor o igual que el índice final")

                        sub_df = df.iloc[inicio: fin+1]
                        for i, row in sub_df.iterrows():
                            date = row['fecha_ingreso']
                            rol = row['Rol']
                            address = row['residencia']
                            rut = row['rut']
                            full_name = row['nombre_completo']
                            nationality = row['nacionalidad']
                            birth_date = row['fecha_de_nacimiento']
                            profession = row['profesion']
                            salary = row['Sueldo']

                            example_contract(date, rol, address, rut, full_name, nationality, birth_date, profession, str(salary))
                            print(f"Contrato generado para {full_name}.")

                    elif sub_opcion == 3:
                        break
                    else:
                        print("Opción no válida. Intente nuevamente.")
                except ValueError:
                    print("Entrada no válida. Por favor, ingrese un número.")
                except IndexError as e:
                    print(e)
        elif opcion == 2:
            while True:
                menu_graficos()
                try:
                    grafico_opcion = int(input("Ingrese el número de la opción deseada: "))
                    if grafico_opcion == 4:
                        break
                    elif grafico_opcion in opciones_graficos:
                        opciones_graficos[grafico_opcion](df)
                    else:
                        print("Opción no válida. Intente nuevamente.")
                except ValueError:
                    print("Entrada no válida. Por favor, ingrese un número.")
        elif opcion == 3:
            print("Gracias por preferirnos")
            break
        else:
            print("Opción no válida. Intente nuevamente.")
    except ValueError:
        print("Entrada no válida. Por favor, ingrese un número.")

