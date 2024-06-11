import sqlite3
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, Mm

def obtener_datos_desde_bd(bd):
    # Se conecta con la base de datos bd
    conexion = sqlite3.connect(bd)
    # Consulta en SQLite
    consulta = '''                   
    SELECT personas.*, Salarios.*
    FROM personas
    INNER JOIN Salarios ON personas.id_rol = id_Salarios
    '''
    df = pd.read_sql_query(consulta, conexion)  # Se recuperan los datos en un DataFrame

    conexion.close()  # Se cierra la conexión

    return df  # Se devuelve el DataFrame

# Función que genera un contrato de ejemplo
def example_contract(date: str, rol: str, address: str, rut: str, full_name: str, nationality: str, birth_date: str, profession: str, salary: str):
    document = Document()
    # Configuración del documento
    font = document.styles['Normal'].font
    font.name = 'Book Antiqua'
    sections = document.sections
    for section in sections:
        section.page_height = Cm(35.56)
        section.page_width = Cm(21.59)
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
        section.header_distance = Mm(12.7)
        section.footer_distance = Mm(12.7)
    # Encabezado
    header = document.sections[0].header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture("imagenes/header.png")
    document.add_picture('imagenes/logo.png')
    h = document.add_paragraph('')
    h.add_run('CONTRATO DE PRESTACIÓN DE SERVICIOS A HONORARIOS\n').bold = True
    h.alignment = 1
    font.size = Pt(10)
    h.add_run(
        '............................................................................................................................................................................................................................. ').bold = True
    # Cuerpo del contrato
    p = document.add_paragraph(f'En Temuco, a {str(date)}, entre la Corporación de Innovación y Desarrollo Tecnológico, Rut 78.898.766-4, representada por su Director General don(a) Roberto Gomez Bolainas, Cédula de Identidad Nº 10.678.990-2, ambos domiciliados en Caupolican 455 de esta ciudad, en adelante la “Corporación” y  {full_name}, de nacionalidad {nationality}, de profesión {profession}, nacido el {birth_date}, con domicilio en {address}, Cédula de Identidad N° {rut}, en adelante, el “Prestador de Servicios”, se ha convenido el siguiente contrato de prestación de servicios a honorarios: \n')
    font.size = Pt(8)
    p1 = document.add_paragraph('')
    p1.add_run('PRIMERO        :').bold = True
    p1.add_run('En el marco del acuerdo de servicios profesionales fechado el 11 de noviembre de 2020, establecido entre la Agencia Nacional de Estándares Educativos y la Corporación de Innovación y Desarrollo Tecnológico , y ratificado según la Resolución Exenta N°603 del 23 de noviembre de 2020, la Corporación encarga los servicios profesionales del Prestador de Servicios, para que ejecute la siguiente tarea en el proyecto "Evaluación de competencias específicas y metodologías de aprendizaje artificial 2020, ID 67703-20-JJ90."')
    p1.add_run('\n SEGUNDO        :').bold = True
    p1.add_run('El rol a desempeñar es de '+rol+'.')
    p1.add_run('\n TERCERO        :').bold = True
    p1.add_run('El plazo para la realización de la prestación de servicios encomendada será el '+str(date))
    p1.add_run('\n CUARTO        :').bold = True
    p1.add_run('Por el servicio profesional efectivamente realizado, se pagara un monto bruto variable, el cual corresponderá a cada rol dentro de la empresa capacitación, de acuerdo al siguiente detalle: ')
    table = document.add_table(rows=2, cols=2)
    table.alignment = 1
    hdr_cells0 = table.rows[0].cells
    hdr_cells0[0].text = 'Rol'
    hdr_cells0[1].text = 'Monto Bruto'
    hdr_cells = table.rows[1].cells
    hdr_cells[0].text = rol
    hdr_cells[1].text = salary
    p1.add_run('\n QUINTO        :').bold = True
    p1.add_run('El Prestador de Servicios acepta el encargo y las condiciones precedentes.')
    p1.add_run('\n SEXTO        :').bold = True
    p1.add_run('El Prestador de Servicios está obligado a mantener la confidencialidad de todos los materiales utilizados, conforme al Acuerdo de Confidencialidad previamente establecido.')
    p1.add_run('\n En comprobante, previa lectura y ratificación, las partes firman.  ').bold = True
    table = document.add_table(rows=2, cols=2)
    table.alignment = 1
    hdr_cells0 = table.rows[0].cells[1].add_paragraph()
    r = hdr_cells0.add_run()
    r.add_picture('imagenes/firma.png')
    hdr_cells = table.rows[1].cells
    hdr_cells[0].text = '-----------------------------------------------------------\nEL PRESTADOR DE SERVICIOS'
    hdr_cells[1].text = '-----------------------------------------------------------\np. LA CORPORACION'
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    run = paragraph.add_run('Caupolican 0455, Temuco, Chile, www.corpoindet.cl')
    run.add_picture("imagenes/footer1.png")
    document.save(f'{full_name}.docx')

bd = 'db_personas.db'  # Ruta de la base de datos

df = obtener_datos_desde_bd(bd)  # Se obtienen los datos de la base de datos

print(df)  # Se muestra el DataFrame

contratos = []

while True:
    try:
        inicio = int(input("Ingrese el índice de inicio para generar los contratos: "))
        fin = int(input("Ingrese el índice final para generar los contratos: "))

        if (inicio < 0 or inicio >= len(df)) or (fin < 0 or fin >= len(df)):
            raise IndexError("Índice fuera de rango")
        
        # Verificar que el rango es válido
        if inicio > fin:
            raise ValueError("El índice de inicio debe ser menor o igual que el índice final")
        
        # Se obtienen los datos de las personas seleccionadas
        sub_df = df.iloc[inicio: fin+1]
        for i, row in sub_df.iterrows():
            date = row['fecha_ingreso']
            rol = row['Rol']
            address = row['residencia']
            rut = row['rut']
            full_name = row['nombre_completo']
            nationality = row['nacionalidad']
            birth_date = row['fecha_de_nacimiento']
            profession = row['profesion']
            salary = row['Sueldo']

            # Se genera el contrato para la persona especificada
            example_contract(date, rol, address, rut, full_name, nationality, birth_date, profession, str(salary))
            contratos.append(full_name)
            print(f"Contrato generado para {full_name}.")
    except IndexError:
        print("Intente nuevamente con índices dentro del rango válido.")
    except ValueError:
        print("Ingrese un número válido y asegúrese de que el índice de inicio sea menor o igual que el índice final.")

print("Proceso de generación de contratos finalizado.")
